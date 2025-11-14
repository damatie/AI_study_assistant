"""Helpers for DOC and DOCX study materials."""

from __future__ import annotations

import io
import math
import zipfile
from xml.etree import ElementTree

from docx import Document  # type: ignore
import olefile  # type: ignore

_DOCX_PAGES_XPATH = "{http://schemas.openxmlformats.org/officeDocument/2006/extended-properties}Pages"


def _coalesce_page_count(value: int | str | None) -> int | None:
    if value is None:
        return None
    try:
        pages = int(value)
    except (TypeError, ValueError):
        return None
    return pages if pages > 0 else None


def _estimate_pages_from_words(words: int | None) -> int:
    if not words or words <= 0:
        return 1
    return max(1, math.ceil(words / 400))


def _docx_word_count(document: Document) -> int:
    words = 0
    for paragraph in document.paragraphs:
        if paragraph.text:
            words += len(paragraph.text.split())
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    words += len(cell.text.split())
    return words


def get_docx_page_count(docx_bytes: bytes) -> int:
    """Return the best-effort page count for a DOCX payload."""

    with io.BytesIO(docx_bytes) as buf:
        with zipfile.ZipFile(buf) as archive:
            try:
                with archive.open("docProps/app.xml") as props:
                    tree = ElementTree.parse(props)
                    pages = tree.findtext(_DOCX_PAGES_XPATH)
                    resolved = _coalesce_page_count(pages)
                    if resolved:
                        return resolved
            except Exception:
                pass

    try:
        document = Document(io.BytesIO(docx_bytes))
    except Exception:
        return 1

    word_count = _docx_word_count(document)
    return _estimate_pages_from_words(word_count)


def get_doc_page_count(doc_bytes: bytes) -> int:
    """Return the best-effort page count for a legacy DOC payload."""

    try:
        with olefile.OleFileIO(io.BytesIO(doc_bytes)) as ole:
            metadata = ole.get_metadata()
    except Exception:
        metadata = None

    if metadata:
        resolved = _coalesce_page_count(getattr(metadata, "num_pages", None))
        if resolved:
            return resolved
        word_count = getattr(metadata, "word_count", None)
        if isinstance(word_count, int) and word_count > 0:
            return _estimate_pages_from_words(word_count)

    return 1


def get_office_page_count(file_bytes: bytes, extension: str) -> int:
    """Return a page count for a DOC/DOCX file based on its extension."""

    ext = extension.lower()
    if ext == ".docx":
        return get_docx_page_count(file_bytes)
    if ext == ".doc":
        return get_doc_page_count(file_bytes)
    raise ValueError(f"Unsupported Office extension: {extension}")


def extract_docx_text(docx_bytes: bytes) -> str:
    """Return plain text content from a DOCX payload for fallback prompts."""

    try:
        document = Document(io.BytesIO(docx_bytes))
    except Exception:
        return ""

    segments: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            segments.append(text)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    segments.append(text)
    return "\n\n".join(segments)
